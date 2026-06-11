"""GitHub-style .docx renderer for the revision-diff model.

Requires python-docx (the ``docx`` optional extra). The CLI checks for
the dependency before importing this module, so module-level docx
imports are safe here.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from gdoc.diffrender import select_visible, short_time, split_comments
from gdoc.revdiff import DEFAULT_CONTEXT, hunk_side_text

# GitHub-ish palette
INK = RGBColor(0x24, 0x29, 0x2F)
CTX = RGBColor(0x57, 0x60, 0x6A)
NAVY = RGBColor(0x0B, 0x2B, 0x52)
GREEN_TX = RGBColor(0x0A, 0x5A, 0x24)
RED_TX = RGBColor(0x8B, 0x14, 0x1B)
COLLAPSE = RGBColor(0x8C, 0x95, 0x9D)
HR_COLOR = RGBColor(0xD0, 0xD7, 0xDE)
GREEN_BG = "CCFFD8"
RED_BG = "FFD7D5"
GREEN_BAR = "2DA44E"
RED_BAR = "CF222E"
BLUE_BAR = "0969DA"

# (fill, bar) pairs cycled per comment author
_AUTHOR_STYLES = [
    ("FFF8C5", "D4A72C"),
    ("DDF4FF", "54AEFF"),
    ("FBEFFF", "8250DF"),
    ("FFF1E5", "FB8500"),
]


def _shade_run(run, fill):
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def _left_bar(par, color, sz=20, space=10):
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pbdr.append(left)
    ppr.append(pbdr)


def _cell_shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _cell_box(cell, fill, bar):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "2")
        e.set(qn("w:color"), fill)
        borders.append(e)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), bar)
    borders.append(left)
    tcpr.append(borders)
    margins = OxmlElement("w:tcMar")
    for edge, width in (
        ("top", "60"), ("bottom", "60"), ("left", "160"), ("right", "120"),
    ):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), width)
        e.set(qn("w:type"), "dxa")
        margins.append(e)
    tcpr.append(margins)


def _add_run(
    par, text, color=INK, bg=None, strike=False, bold=False,
    italic=False, size=None,
):
    run = par.add_run(text)
    run.font.color.rgb = color
    if bg:
        _shade_run(run, bg)
    run.font.strike = strike
    run.font.bold = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    return run


def _add_diff_runs(par, runs):
    for r in runs:
        if r["op"] == "equal":
            _add_run(par, r["text"], color=INK)
        elif r["op"] == "del":
            _add_run(par, r["text"], color=RED_TX, bg=RED_BG, strike=True)
        else:
            _add_run(par, r["text"], color=GREEN_TX, bg=GREEN_BG)


class _DocxBuilder:
    def __init__(self, model, context):
        self.model = model
        self.context = context
        self.doc = Document()
        self.author_styles: dict[str, tuple[str, str]] = {}

        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = INK
        for section in self.doc.sections:
            section.left_margin = section.right_margin = Inches(0.9)
            section.top_margin = section.bottom_margin = Inches(0.8)

    def _author_style(self, author):
        if author not in self.author_styles:
            self.author_styles[author] = _AUTHOR_STYLES[
                len(self.author_styles) % len(_AUTHOR_STYLES)
            ]
        return self.author_styles[author]

    def _heading(self, text, level=2, color=NAVY, space_before=14):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(4)
        sizes = {1: 17, 2: 14, 3: 12, 4: 11}
        _add_run(p, text, color=color, bold=True, size=sizes.get(level, 12))

    def _hr(self):
        p = self.doc.add_paragraph()
        _add_run(p, "_" * 78, color=HR_COLOR, size=9)

    def _collapse(self, gap):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        label = "paragraph" if gap == 1 else "paragraphs"
        _add_run(
            p, f"⋯  {gap} unchanged {label}  ⋯",
            color=COLLAPSE, italic=True, size=9,
        )

    def _comment_box(self, c):
        fill, bar = self._author_style(c["author"])
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        tblw = OxmlElement("w:tblW")
        tblw.set(qn("w:type"), "dxa")
        tblw.set(qn("w:w"), "9400")
        table._tbl.tblPr.append(tblw)
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        _cell_shade(cell, fill)
        _cell_box(cell, fill, bar)

        head = cell.paragraphs[0]
        head.paragraph_format.space_after = Pt(2)
        _add_run(head, "💬 " + c["author"], bold=True, size=9.5)
        _add_run(
            head, "   ·   " + short_time(c["createdTime"]),
            size=8.5, color=CTX,
        )
        if c.get("resolved"):
            _add_run(head, "   (resolved)", size=8.5, color=CTX, italic=True)
        if c.get("quoted"):
            quoted = c["quoted"]
            if len(quoted) > 90:
                quoted = quoted[:90] + "…"
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, "on: “" + quoted + "”", italic=True, size=8.5,
                     color=CTX)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, c["content"], size=9.5)
        for r in c.get("replies", []):
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(1)
            _add_run(p, "↳ " + r["author"] + ": ", bold=True, size=9)
            _add_run(p, r["content"], size=9)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _render_hunk(self, hunk):
        kind = hunk["kind"]
        if hunk["block_type"] == "heading":
            if kind == "replace":
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                _add_diff_runs(p, hunk["runs"])
            else:
                side = "old" if kind == "delete" else "new"
                color = {
                    "insert": GREEN_TX, "delete": RED_TX,
                }.get(kind, NAVY)
                self._heading(
                    hunk_side_text(hunk, side),
                    level=hunk.get("level", 2), color=color,
                )
            return

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        bullet = hunk["block_type"] == "listitem"
        prefix = "•  " if bullet else ""
        if bullet:
            p.paragraph_format.left_indent = Inches(0.3)

        if kind == "equal":
            if prefix:
                _add_run(p, prefix, color=CTX)
            _add_run(p, hunk_side_text(hunk, "new"), color=CTX)
        elif kind == "insert":
            _left_bar(p, GREEN_BAR)
            p.paragraph_format.left_indent = Inches(0.3 if bullet else 0.12)
            if prefix:
                _add_run(p, prefix, color=GREEN_TX, bg=GREEN_BG)
            _add_run(
                p, hunk_side_text(hunk, "new"),
                color=GREEN_TX, bg=GREEN_BG,
            )
        elif kind == "delete":
            _left_bar(p, RED_BAR)
            p.paragraph_format.left_indent = Inches(0.3 if bullet else 0.12)
            if prefix:
                _add_run(p, prefix, color=RED_TX, bg=RED_BG, strike=True)
            _add_run(
                p, hunk_side_text(hunk, "old"),
                color=RED_TX, bg=RED_BG, strike=True,
            )
        else:
            _left_bar(p, BLUE_BAR)
            p.paragraph_format.left_indent = Inches(0.3 if bullet else 0.12)
            if prefix:
                _add_run(p, prefix, color=INK)
            _add_diff_runs(p, hunk["runs"])

    def build(self, out_path):
        model = self.model
        hunks = model["hunks"]
        comments = model.get("comments", [])
        by_hunk, appendix = split_comments(comments)
        keep = select_visible(
            hunks, self.context, comment_hunks=set(by_hunk),
        )

        title = self.doc.add_paragraph()
        title.paragraph_format.space_after = Pt(2)
        _add_run(
            title, f"{model['doc']['name']} — revision diff",
            bold=True, size=22, color=NAVY,
        )
        meta = self.doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(10)
        old, new = model["old"], model["new"]
        _add_run(
            meta, f"rev {old['id']}  ({short_time(old['modifiedTime'])})",
            color=RED_TX, size=10,
        )
        _add_run(meta, "   →   ", color=CTX, size=10)
        _add_run(
            meta, f"rev {new['id']}  ({short_time(new['modifiedTime'])})",
            color=GREEN_TX, size=10,
        )
        legend = self.doc.add_paragraph()
        legend.paragraph_format.space_after = Pt(2)
        _add_run(legend, " added ", color=GREEN_TX, bg=GREEN_BG, size=9.5)
        _add_run(legend, "    ", size=9.5)
        _add_run(
            legend, " removed ", color=RED_TX, bg=RED_BG,
            strike=True, size=9.5,
        )
        _add_run(legend, "    │ blue bar = reworded line", color=CTX,
                 size=9.5)
        if comments:
            _add_run(
                legend, "   │ 💬 comment threads (color-coded by author)",
                color=CTX, size=9.5,
            )
        self._hr()

        gap = 0
        for i, hunk in enumerate(hunks):
            if not keep[i]:
                gap += 1
                continue
            if gap:
                self._collapse(gap)
                gap = 0
            self._render_hunk(hunk)
            for c in by_hunk.get(i, []):
                self._comment_box(c)
        if gap:
            self._collapse(gap)

        if appendix:
            self._hr()
            self._heading(
                "Other comment threads (anchored outside shown hunks)",
                level=2,
            )
            for c in appendix:
                self._comment_box(c)

        self.doc.save(out_path)
        return {
            "hunks": len(hunks),
            "shown": sum(keep),
            "comments_inline": len(comments) - len(appendix),
            "comments_appendix": len(appendix),
        }


def render_docx(
    model: dict, out_path: str, context: int = DEFAULT_CONTEXT,
) -> dict:
    """Render the diff model to a styled .docx file. Returns stats."""
    return _DocxBuilder(model, context).build(out_path)
